import json
import os
from docx import Document
from docx2pdf import convert
from docx.shared import Inches

def generar_constancia(datos, nombreDoc, idUsuario, carpeta_docs, carpeta_firmas, datos2=None):
    # --- 1. CONFIGURACIÓN DE RUTAS ---
    ruta_docx = os.path.join(carpeta_docs, f"{nombreDoc}.docx")
    salida_docx = os.path.join(carpeta_docs, f"temp_{nombreDoc}.docx")
    salida_pdf = os.path.join(carpeta_docs, f"{nombreDoc}.pdf")
    
    # Ruta de la firma del usuario actual (el que está logueado)
    ruta_firma_usuario = os.path.join(carpeta_firmas, f"firma_{idUsuario}.png")

    # Validar que exista la plantilla Word
    if not os.path.exists(ruta_docx):
        print(f"❌ Error: No existe la plantilla en {ruta_docx}")
        return None

    # --- 2. LIMPIEZA DE DATOS ---
    if isinstance(datos, str):
        try:
            datos = json.loads(datos)
        except:
            # Intento de rescate si el formato es texto plano antiguo
            dic = {}
            partes = datos.split(";")
            for p in partes:
                if ":" in p:
                    k, v = p.split(":", 1)
                    dic[k.strip()] = v.strip()
            datos = dic

    try:
        doc = Document(ruta_docx)

        # --- FUNCIÓN MAESTRA DE REEMPLAZO ---
        def procesar_parrafo(parrafo):
            texto_actual = parrafo.text

            # A. BUSCAR FIRMAS DINÁMICAS (Desde el JSON)
            # Ejemplo: JSON dice "VAR_FIRMA_JEFE": "jefe_sistemas.png"
            for k, v in datos.items():
                # Si la variable empieza con "VAR_FIRMA_" y está en el documento
                if k.startswith("VAR_FIRMA_") and k in texto_actual:
                    nombre_imagen = str(v).strip() # El valor del JSON es el nombre del archivo
                    ruta_imagen_extra = os.path.join(carpeta_firmas, nombre_imagen)

                    # Borramos el texto de la variable (ej: VAR_FIRMA_JEFE)
                    parrafo.text = parrafo.text.replace(k, "")
                    
                    if os.path.exists(ruta_imagen_extra):
                        # Insertamos la imagen
                        run = parrafo.add_run()
                        run.add_picture(ruta_imagen_extra, width=Inches(1.5))
                        print(f"🖼️ Firma dinámica insertada: {k} -> {nombre_imagen}")
                    else:
                        print(f"⚠️ Se pidió la firma '{k}' pero no existe el archivo: {ruta_imagen_extra}")
                        # Opcional: Agregar texto de error en el PDF
                        # parrafo.add_run("[Falta Imagen]")

            # B. REEMPLAZO DE TEXTO NORMAL
            for k, v in datos.items():
                # Ignoramos las llaves de firma que ya procesamos arriba
                if not k.startswith("VAR_FIRMA_") and k in parrafo.text:
                    texto_nuevo = parrafo.text.replace(k, str(v))
                    # Limpiamos runs para evitar problemas de formato
                    for run in parrafo.runs: run.text = ""
                    parrafo.runs[0].text = texto_nuevo

            # C. FIRMA DEL USUARIO (VAR_FIRMA)
            # Esta es la firma de la persona logueada (idUsuario)
            if "VAR_FIRMA" in parrafo.text:
                if os.path.exists(ruta_firma_usuario):
                    parrafo.text = parrafo.text.replace("VAR_FIRMA", "")
                    run = parrafo.add_run()
                    run.add_picture(ruta_firma_usuario, width=Inches(1.5))
                    print(f"🖼️ Firma usuario insertada (ID {idUsuario})")
                else:
                    # Si no tiene firma, solo borramos la variable
                    parrafo.text = parrafo.text.replace("VAR_FIRMA", "")
                    print(f"⚠️ El usuario {idUsuario} no tiene firma registrada.")

        # --- 3. EJECUTAR REEMPLAZO EN TODO EL DOCUMENTO ---
        
        # 3.1 Párrafos del cuerpo principal
        for p in doc.paragraphs:
            procesar_parrafo(p)

        # 3.2 Tablas (Soporte para tablas anidadas)
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for p in celda.paragraphs:
                        procesar_parrafo(p)
                    # Tablas dentro de tablas
                    for subtabla in celda.tables:
                        for subfila in subtabla.rows:
                            for subcelda in subfila.cells:
                                for sub_p in subcelda.paragraphs:
                                    procesar_parrafo(sub_p)

        # --- 4. GUARDAR Y CONVERTIR ---
        doc.save(salida_docx)
        convert(salida_docx, salida_pdf)

        if os.path.exists(salida_docx):
            os.remove(salida_docx) # Borrar temporal

        print(f"✅ PDF generado correctamente: {salida_pdf}")
        return salida_pdf

    except Exception as e:
        print(f"❌ Error crítico en conversión: {e}")
        return None